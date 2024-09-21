import time
import docx
from docx import Document
import fitz
import pypandoc
from pdfminer.high_level import extract_text
import os
from PyPDF2 import PdfReader
import shutil
from PyPDF2 import PdfReader, PdfWriter
from fpdf import FPDF

def type_like_effect(text, delay=0.1):
    for char in text:
        print(char, end='', flush=True)
        time.sleep(delay)
def get_docx_metadata(file):
    try:
        doc = docx.Document(file.name)
        core_properties = doc.core_properties
        metadata = {
            'Title': core_properties.title,
            'Author': core_properties.author,
            'Last Modified By': core_properties.last_modified_by,
            'Created': core_properties.created,
            'Last Modified': core_properties.modified,
            'Revision': core_properties.revision,
            'Keywords': core_properties.keywords,
        }
        return metadata
    except Exception as e:
        return {"Error": f"Could not extract DOCX metadata: {str(e)}"}

def replace(file, file_name, file_extension, word_to_replace):
    file.seek(0)
    file = file.read()
    if word_to_replace not in file:
        type_like_effect('Word not found in file.\n', delay=0.02)
        time.sleep(1)
        fifth_option(file, file_name, file_extension)
    else:
        type_like_effect('Enter word/character to replace with: ', delay=0.02)
        word_to_replace_with = input()
        type_like_effect('Would you like to replace all occurrences of the word/character? (y/n): ', delay=0.02)
        replace_all = input()
        if replace_all.lower() == 'y':
            filee = file.replace(word_to_replace, word_to_replace_with)
            type_like_effect('Replacing all occurrences of the word/character...\n', delay=0.02)
            time.sleep(1)
            with open(file_name, 'w') as file:
                file.write(filee)
            type_like_effect('Replacement successful!\n', delay=0.002)
            time.sleep(1)
        elif replace_all.lower() == 'n':
            lines_with_word = [i for i, line in enumerate(file.split('\n')) if word_to_replace in line]
            if not lines_with_word:
                type_like_effect('Word not found in file.\n', delay=0.02)
                time.sleep(1)
                fifth_option(file, file_name, file_extension)
            else:
                type_like_effect('Enter the line number to replace: ', delay=0.02)
                try:
                    line_number = int(input())
                except ValueError:
                    type_like_effect('Invalid input. Please enter a valid line number.\n', delay=0.02)
                    time.sleep(1)
                    fifth_option(file, file_name, file_extension)
                if line_number < 1 or line_number > len(lines_with_word):
                    type_like_effect('Invalid line number.\n', delay=0.02)
                    time.sleep(1)
                    fifth_option(file, file_name, file_extension)
                elif line_number in lines_with_word:
                    if word_to_replace in file.split('\n')[line_number - 1].count(word_to_replace) > 1:
                        type_like_effect('Enter the occurrence number to replace (1 for first occurrence, 2 for second occurrence, etc.): ', delay=0.02)
                        try:
                            occurrence_number = int(input())
                        except ValueError:
                            type_like_effect('Invalid input. Please enter a valid occurrence number.\n', delay=0.02)
                            time.sleep(1)
                            fifth_option(file, file_name, file_extension)
                        if occurrence_number < 1 or occurrence_number > file.split('\n')[line_number - 1].count(word_to_replace):
                            type_like_effect('Invalid occurrence number.\n', delay=0.02)
                            time.sleep(1)
                            fifth_option(file, file_name, file_extension)
                        else:
                            file = file.replace(word_to_replace, word_to_replace_with, occurrence_number)
                            type_like_effect('Replacing the occurrence...\n', delay=0.02)
                            time.sleep(1)
                            with open(file_name, 'w') as file:
                                file.write(file)
                            type_like_effect('Replacement successful!\n', delay=0.002)
                            time.sleep(1)
                    else:
                        file = file.replace(word_to_replace, word_to_replace_with, 1)
                        type_like_effect('Replacing the occurrence...\n', delay=0.02)
                        time.sleep(1)
                        with open(file_name, 'w') as file:
                            file.write(file)
                        type_like_effect('Replacement successful!\n', delay=0.002)
                        time.sleep(1)
                else:
                    type_like_effect('Invalid line number.\n', delay=0.02)
                    time.sleep(1)
                    fifth_option(file, file_name, file_extension)
        else:
            type_like_effect('Invalid input. Please enter "y" or "n".\n', delay=0.02)
            time.sleep(1)
            fifth_option(file, file_name, file_extension) 

def get_pdf_metadata(file):
    """Extract metadata from a PDF file."""
    try:
        reader = PdfReader(file.name)
        metadata = reader.metadata
        if metadata:
            return {
                'Title': metadata.title,
                'Author': metadata.author,
                'Subject': metadata.subject,
                'Producer': metadata.producer,
                'Creation Date': metadata.creation_date,
                'Modification Date': metadata.mod_date,
            }
        else:
            return {"Error": "No metadata found in the PDF file."}
    except Exception as e:
        return {"Error": f"Could not extract PDF metadata: {str(e)}"}
def end():
    type_like_effect('Would you like to do anything else? (y/n): ', delay = 0.02)
    answer = input()
    if answer.lower() == 'y':
        pass
    elif answer.lower() == 'n':
        type_like_effect("Thank you for using Sanmi's program", delay = 0.02)
        time.sleep(1)
        type_like_effect("\nGoodbye🤗😘")
        time.sleep(1)
        exit()
    else:
        type_like_effect('Invalid option...\nTry again\n', delay = 0.01)
        end()

def pdf_to_txt(pdf_file, output_txt_file):
    reader = PdfReader(pdf_file)
    with open(output_txt_file, 'w', encoding='utf-8') as txt_file:
        for page in reader.pages:
            text = page.extract_text()
            if text:
                txt_file.write(text)

def txt_to_pdf(txt_file, output_pdf_file):
    """
    Converts a TXT file to a PDF file using UTF-8 encoding.
    """
    try:
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=12)

        with open(txt_file, 'r', encoding='utf-8') as txt:
            for line in txt:
                # Ensure that non-ASCII characters are handled
                pdf.multi_cell(0, 10, line.encode('latin-1', 'replace').decode('latin-1'))

        pdf.output(output_pdf_file)
        print(f"Successfully converted {txt_file} to {output_pdf_file}")
    except Exception as e:
        print(f"Error converting TXT to PDF: {e}")

def docx_to_txt(docx_file, output_txt_file):
    doc = Document(docx_file)
    with open(output_txt_file, 'w', encoding='utf-8') as txt_file:
        for para in doc.paragraphs:
            txt_file.write(para.text + '\n')

def txt_to_docx(txt_file, output_docx_file):
    doc = Document()
    with open(txt_file, 'r', encoding='utf-8') as txt:
        for line in txt:
            doc.add_paragraph(line)
    doc.save(output_docx_file)

def docx_to_pdf(docx_file, output_pdf_file):
    temp_txt_file = "temp_text.txt"
    docx_to_txt(docx_file, temp_txt_file)
    txt_to_pdf(temp_txt_file, output_pdf_file)
    os.remove(temp_txt_file)  

def pdf_to_docx(pdf_file, output_docx_file):
    reader = PdfReader(pdf_file)
    doc = Document()
    for page in reader.pages:
        text = page.extract_text()
        if text:
            doc.add_paragraph(text)
    doc.save(output_docx_file)

def exit_option():
    while True:
        type_like_effect('\nExiting to main menu...\n', delay=0.02)
        time.sleep(1)
        time.sleep(1)
        break

def get_file_type(file_path):
    with open(file_path, 'rb') as file:
        file_signature = file.read(4)
    if file_signature.startswith(b'%PDF'):
        return "PDF File"
    elif file_signature == b'\x50\x4B\x03\x04':  
        return "DOCX File"
    elif file_signature == b'\xd0\xcf\x11\xe0':
        return 'DOC File'
    else:
        with open(file_path, 'rb') as file:
            content = file.read(1024)
            try:
                content.decode('utf-8')
                return "Text File"
            except UnicodeDecodeError:
                return "Unknown File Type"
def read_docx(file):
    doc = docx.Document(file)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)
def read_pdf(file):
    doc = fitz.open(file)
    text = []
    for page_num in range(doc.page_count):
        page = doc.load_page(page_num)   
        text.append(page.get_text())   
        return '\n'.join(text)
def convert_doc_to_docx(input_doc):
    output_docx = input_doc.replace('.doc', '.docx')
    pypandoc.convert_file(input_doc, 'docx', outputfile=output_docx)
    return output_docx
def menu():
    type_like_effect(f'''
----------------------------Menu----------------------------
1-->Read file<--
2-->Total number of characters/words<--
3-->Total number of lines<--
4-->Find character/word/pattern<--
5-->Replace character/word<--
6-->Copy file<--
7-->Get file info<--
8-->Convert file<--
9-->Delete file<--
10-->Open new file<--
11-->Options<--
12-->Exit<--
Type 'dir' to see all files in the current directory.
---------------------------------------------------------------\n''', delay = 0.002)


    

def second_option(file, file_extension):
        type_like_effect(f'Total number of characters or words(c/w): ', delay=0.02)
        user_input = input()
        if file_extension == 'Text File':
            file.seek(0)
            if user_input.lower() == 'c':
                type_like_effect(f'Total number of characters: {len(file.read())}\n', delay=0.02)
            elif user_input.lower() == 'w':
                type_like_effect(f'Total number of words: {len(file.read().split())}\n', delay=0.02)
            else:
                type_like_effect('Invalid input.\n', delay=0.02)
                second_option(file, file_extension)
        else:
            if user_input.lower() == 'c':
                type_like_effect(f'Total number of characters: {len(file)}\n', delay=0.02)
            elif user_input.lower() == 'w':
                type_like_effect(f'Total number of words: {len(file.split())}\n', delay=0.02)
            else:
                type_like_effect('Invalid input.\n', delay=0.02)
                second_option(file, file_extension)

def third_option(file, file_extension):
    if file_extension == 'Text File':
        file.seek(0)
        type_like_effect(f'Total number of lines: {len(file.readlines())}\n', delay=0.02)
    else:
        type_like_effect(f'Total number of lines: {len(file.splitlines())}\n', delay=0.02)

def fourth_option_menu():
        type_like_effect('''
-----------------Please select an option-----------------
1--Display file content
2--Total occurrence of selected word/character
3--Display lines containing selected word/character
4--Display lines not containing selected word/character
5--Menu
6--Exit
''', delay = 0.02)
def fourth_option(file, file_extension):
    while True:
        type_like_effect('Enter option (5 for menu): ', delay=0.02)
        user_input = input()
        if user_input == '1':
            if file_extension == 'Text File':
                file.seek(0)
                type_like_effect(f'{file.read()}\n', delay=0.02)
            else:
                type_like_effect(f'{file}\n', delay=0.02)

        elif user_input == '2':
            type_like_effect('Enter the word/character to search for: ', delay=0.02)
            search_word = input()
            if file_extension == 'Text File':
                file.seek(0)
                file_content = file.read()
                if search_word not in file_content:
                    type_like_effect(f'The word "{search_word}" was not found in the file.\n', delay=0.02)
                else:
                    try:
                        occurrences = file_content.count(search_word)
                        type_like_effect(f'Total occurrence of "{search_word}": {occurrences}\n', delay=0.02)
                    except Exception as e:
                        type_like_effect(f'Error: {e}\n', delay=0.02)
            else:
                occurrences = file.count(search_word)
                if search_word not in file:
                    type_like_effect(f'The word "{search_word}" was not found in the file.\n', delay=0.02)
                else:
                    type_like_effect(f'Total occurrence of "{search_word}": {occurrences}\n', delay=0.02)

        elif user_input == '3':
            type_like_effect('Enter the word/character to search for: ', delay=0.02)
            search_word = input()
            if file_extension == 'Text File':
                file.seek(0)
                lines = file.readlines()
            else:
                lines = file.splitlines()
            matching_lines = [line for line in lines if search_word in line]
            if matching_lines:
                type_like_effect(f'Lines containing "{search_word}":\n', delay=0.02)
                for line in matching_lines:
                    type_like_effect(line + '\n', delay=0.02)
            else:
                type_like_effect(f'No lines containing "{search_word}" found.\n', delay=0.02)

        elif user_input == '4':
            type_like_effect('Enter the word/character to search for: ', delay=0.02)
            search_word = input()
            if file_extension == 'Text File':
                file.seek(0)
                lines = file.readlines()
            else:
                lines = file.splitlines()
            non_matching_lines = [line for line in lines if search_word not in line]
            if non_matching_lines:
                type_like_effect(f'Lines not containing "{search_word}":\n', delay=0.02)
                for line in non_matching_lines:
                    type_like_effect(line + '\n', delay=0.02)
            else:
                type_like_effect(f'All lines contain "{search_word}".\n', delay=0.02)

        elif user_input == '5':
            fourth_option_menu()


        elif user_input == '6':
            type_like_effect('Exiting...', delay=0.02)
            time.sleep(1)
            type_like_effect('Thank you for using our file search tool!', delay=0.02)
            time.sleep(1)
            time.sleep(1)
            exit_option()
            break

        else:
            type_like_effect('Invalid option. Please try again.', delay=0.02)

def seventh_option(file_name):   
    type_like_effect('Getting metadata...\n', delay=0.02)
    time.sleep(1)
    type_like_effect('Metadata retrieved successfully!\n', delay=0.02)
    time.sleep(1)

    type_like_effect(f'File name: {file_name}\n', delay=0.02)
    type_like_effect(f'File size: {os.path.getsize(file_name)} bytes\n', delay=0.02)
    type_like_effect(f'File creation time: {time.ctime(os.path.getctime(file_name))}\n', delay=0.02)
    type_like_effect(f'File last modified time: {time.ctime(os.path.getmtime(file_name))}\n', delay=0.02)
    type_like_effect(f'File type: {file_name.split(".")[-1]}\n', delay=0.02)
    type_like_effect(f'File path: {os.path.abspath(file_name)}\n', delay=0.02)

    file_extension = file_name.split('.')[-1].lower()
    
    if file_extension == 'DOCX File':
        metadata = get_docx_metadata(file_name)
        type_like_effect(f'DOCX Metadata:\n', delay=0.02)
        for key, value in metadata.items():
            type_like_effect(f'{key}: {value}\n', delay=0.02)

    elif file_extension == 'PDF File':
        metadata = get_pdf_metadata(file_name)
        type_like_effect(f'PDF Metadata:\n', delay=0.02)
        for key, value in metadata.items():
            type_like_effect(f'{key}: {value}\n', delay=0.02)

    else:
        type_like_effect('No additional metadata available for this file type.\n', delay=0.02)

def copy_file(file_name, file_extension, output_dir=None):
    base_name = os.path.basename(file_name)
    global new_file_name
    new_file_name = f"copy_of_{base_name}"
    if output_dir:
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)  
        new_file_name = os.path.join(output_dir, new_file_name)
    else:
        new_file_name = os.path.join(os.path.dirname(file_name), new_file_name)

    if file_extension == 'Text File':
        shutil.copyfile(file_name, new_file_name)
    elif file_extension == 'DOCX File':
        copy_docx_file(file_name, new_file_name)
    elif file_extension == 'PDF File':
        copy_pdf_file(file_name, new_file_name)
    elif file_extension == 'csv':
        shutil.copyfile(file_name, new_file_name)
    else:
        type_like_effect(f"Unsupported file type: {file_extension}", delay=0.02)
        return None

    print(f"File copied successfully to: {new_file_name}")
    return new_file_name

def copy_docx_file(file_name, new_file_name):
    doc = Document(file_name)
    doc.save(new_file_name)

def copy_pdf_file(file_name, new_file_name):
    reader = PdfReader(file_name)
    writer = PdfWriter()

    for page_num in range(len(reader.pages)):
        writer.add_page(reader.pages[page_num])

    with open(new_file_name, "wb") as output_pdf:
        writer.write(output_pdf)
def sixth_option(file_name):
    type_like_effect('Copying file...\n', delay=0.02)
    time.sleep(1)
    copy_file(file_name)
    type_like_effect('File copied successfully!\n', delay=0.02)
    time.sleep(1)
def eighth_option(file_name, file_extension):

    type_like_effect('Are you sure you want to convert this file? (y/n): ', delay=0.02)
    confirmation = input()
    if confirmation.lower() == 'y':
        while True:
            type_like_effect('What file type do you want to convert it to? (docx/pdf/txt): ', delay=0.02)
            conversion_type = input()
            if conversion_type.lower() == 'txt':
                if file_extension == 'Text File':
                    type_like_effect('File is already a text file. No conversion needed.\n', delay=0.02)
                    time.sleep(1)
                    continue    
                elif file_extension == 'DOCX File':
                    type_like_effect('Converting DOCX file to text...\n', delay=0.02)
                    time.sleep(1)
                    docx_to_txt(file_name, file_name.replace('.docx', '.txt'))
                    type_like_effect('Conversion successful!\n', delay=0.02)
                    time.sleep(1)
                    break
                elif file_extension == 'PDF File':
                    type_like_effect('Converting PDF file to text...\n', delay=0.02)
                    time.sleep(1)
                    pdf_to_txt(file_name, file_name.replace('.pdf', '.txt'))
                    type_like_effect('Conversion successful!\n', delay=0.02)
                    time.sleep(1)
                    break
                else:
                    type_like_effect('Unsupported file type.\n', delay=0.02)
                    time.sleep(1)
                    continue
            elif conversion_type.lower() == 'docx':
                if file_extension == 'Text File':
                    type_like_effect('Converting text file to DOCX...\n', delay=0.02)
                    time.sleep(1)
                    txt_to_docx(file_name, file_name.replace('.txt', '.docx'))
                    type_like_effect('Conversion successful!\n', delay=0.02)
                    time.sleep(1)
                    break
                elif file_extension == 'DOCX File':
                    type_like_effect('File is already a DOCX file. No conversion needed.\n', delay=0.02)
                    time.sleep(1)
                    continue
                elif file_extension == 'PDF File':
                    type_like_effect('Converting PDF file to DOCX...\n', delay=0.02)
                    time.sleep(1)
                    pdf_to_docx(file_name, file_name.replace('.pdf', '.docx'))
                    type_like_effect('Conversion successful!\n', delay=0.02)
                    time.sleep(1)
                    break
                else:
                    type_like_effect('Unsupported file type.\n', delay=0.02)
                    time.sleep(1)
                    continue
            elif conversion_type.lower() == 'pdf':
                if file_extension == 'Text File':
                    type_like_effect('Converting text file to PDF...\n', delay=0.02)
                    time.sleep(1)
                    txt_to_pdf(file_name, file_name.replace('.txt', '.pdf'))
                    type_like_effect('Conversion successful!\n', delay=0.02)
                    time.sleep(1)
                    break
                elif file_extension == 'DOCX File':
                    type_like_effect('Converting DOCX file to PDF...\n', delay=0.02)
                    time.sleep(1)
                    docx_to_pdf(file_name, file_name.replace('.docx', '.pdf'))
                    type_like_effect('Conversion successful!\n', delay=0.02)
                    time.sleep(1)
                    break
                elif file_extension == 'PDF File':
                    type_like_effect('File is already a PDF file. No conversion needed.\n', delay=0.02)
                    time.sleep(1)
                    continue
                else:
                    type_like_effect('Unsupported file type.\n', delay=0.02)
                    time.sleep(1)
                    continue
            else:
                    type_like_effect('Unsupported file type.\n', delay=0.02)
                    time.sleep(1)
                    continue
    elif confirmation.lower() == 'n':
        type_like_effect('Conversion cancelled.\n', delay=0.02)
        time.sleep(1)
        exit_option()





def fifth_option(file, file_name, file_extension):
    while True:
        type_like_effect('Enter word/character to replace (enter "dis" to display file or "exit" to main menu): ', delay=0.02)
        word_to_replace = input()
        
        if word_to_replace == 'exit':
            break
        if word_to_replace == 'dis':
            type_like_effect(f'{file}\n', delay=0.02)
            continue
        
        # Handle different file types
        if file_extension == 'Text File':
            replace(file, file_name, file_extension, word_to_replace)
            break
        
        elif file_extension == 'PDF File':
            # Convert PDF to temporary text file, process, and then back to PDF
            pdf_to_txt(file_name, 'temp.txt')
            with open('temp.txt', 'r+') as temp_file:
                replace(temp_file, 'temp.txt', 'Text File', word_to_replace)
            txt_to_pdf('temp.txt', file_name)
            os.remove('temp.txt')  # Make sure temp.txt is removed only after it’s closed
            break   
        
        elif file_extension == 'DOCX File':
            # Convert DOCX to temporary text file, process, and then back to DOCX
            docx_to_txt(file_name, 'temp.txt')
            with open('temp.txt', 'r+') as temp_file:
                replace(temp_file, 'temp.txt', 'Text File', word_to_replace)
            txt_to_docx('temp.txt', file_name)
            os.remove('temp.txt')  # Make sure temp.txt is removed only after it's closed
            break
        else:
            type_like_effect('Unsupported file type.\n', delay=0.02)
            time.sleep(1)
            exit_option()

        